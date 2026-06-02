import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, size=None, bold=False, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_font(paragraph, size=10.5):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def set_paragraph_border(paragraph, color="BFBFBF", size="6", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_hyperlink(paragraph, text, url):
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_doc(data, out_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.16
    normal.paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(data.get("title", "新闻项目简报"))
    set_run_font(r, size=20, bold=True, color=(32, 32, 32))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(data.get("subtitle", ""))
    set_run_font(r, size=10.5, color=(89, 89, 89))

    lead = doc.add_paragraph()
    lead.add_run("导语：").bold = True
    lead.add_run(
        "本简报按日期窗口、来源可信度、领域覆盖和问题意识筛选；每条保留来源标识和原文链接。"
    )
    set_paragraph_font(lead)

    body_title = doc.add_paragraph()
    body_title.paragraph_format.space_before = Pt(8)
    body_title.paragraph_format.space_after = Pt(8)
    r = body_title.add_run("重点动态")
    set_run_font(r, size=14, bold=True, color=(32, 32, 32))
    set_paragraph_border(body_title, color="BFBFBF", size="8", space="3")

    items = data.get("items", [])
    sections = data.get("sections", [])
    index = 1
    for section_name in sections:
        section_items = [item for item in items if item.get("section") == section_name]
        if not section_items:
            continue
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(10)
        sp.paragraph_format.space_after = Pt(4)
        r = sp.add_run(section_name)
        set_run_font(r, size=13, bold=True, color=(32, 32, 32))
        set_paragraph_border(sp, color="BFBFBF", size="6", space="2")

        for item in section_items:
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(7)
            h.paragraph_format.space_after = Pt(2)
            r = h.add_run(f"{index}. {item['title']}")
            set_run_font(r, size=12, bold=True, color=(32, 32, 32))

            meta = doc.add_paragraph()
            meta.paragraph_format.space_after = Pt(2)
            r = meta.add_run(f"领域：{item['field']}    来源：{item['source']}")
            set_run_font(r, size=9.5, color=(89, 89, 89))

            if item.get("viewpoint"):
                vp = doc.add_paragraph()
                vp.paragraph_format.space_after = Pt(2)
                label = "观点/判断："
                if item.get("viewpoint_person"):
                    text = f"{label}{item['viewpoint_person']}认为，{item['viewpoint']}"
                elif item.get("viewpoint_org"):
                    text = f"{label}{item['viewpoint_org']}显示/指出，{item['viewpoint']}"
                else:
                    text = f"{label}{item['viewpoint']}"
                r = vp.add_run(text)
                set_run_font(r, size=9.5, color=(64, 64, 64))

            body = doc.add_paragraph(item["body"])
            body.paragraph_format.first_line_indent = Inches(0.22)
            body.paragraph_format.space_after = Pt(2)
            set_paragraph_font(body)

            src = doc.add_paragraph()
            src.paragraph_format.space_after = Pt(5)
            src.add_run("原文链接：")
            add_hyperlink(src, item["url"], item["url"])
            set_paragraph_font(src, size=9)
            set_paragraph_border(src, color="E7E6E6", size="4", space="4")
            index += 1

    doc.core_properties.title = data.get("title", "新闻项目简报")
    doc.core_properties.subject = "新闻项目简报"
    doc.core_properties.author = "Codex"
    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate a Chinese news brief DOCX from JSON.")
    parser.add_argument("input_json")
    parser.add_argument("--out", required=True)
    args = parser.parse_args()

    data = json.loads(Path(args.input_json).read_text(encoding="utf-8"))
    build_doc(data, args.out)
    print(args.out)


if __name__ == "__main__":
    main()
